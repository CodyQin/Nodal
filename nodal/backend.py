import os
import json
import io
import re
import logging
import time
import tempfile
import shutil
from collections import Counter
from typing import List, Optional, Dict, Any, Generator
from PyPDF2 import PdfReader

from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.responses import StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from dotenv import load_dotenv

from google import genai
from google.genai import types
from docx import Document
import networkx as nx

# -------------------------
# Logging
# -------------------------
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s",
    handlers=[
        logging.FileHandler("nodal_backend.log", encoding="utf-8"),
        logging.StreamHandler(),
    ],
)
logger = logging.getLogger(__name__)

# Prioritize .env.local then .env
if os.path.exists(".env.local"):
    load_dotenv(".env.local")
    logger.info("Loaded environment from .env.local")
else:
    load_dotenv(".env")
    logger.info("Loaded environment from .env")

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# -------------------------
# Gemini client
# -------------------------
api_key = os.getenv("API_KEY") or os.getenv("GEMINI_API_KEY")
if not api_key:
    logger.error("No API key found! Please set API_KEY or GEMINI_API_KEY in .env.local")
    client = None
else:
    client = genai.Client(api_key=api_key)
    logger.info("Gemini client initialized.")


# -------------------------
# Helper: extract text
# -------------------------
def extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        doc = Document(io.BytesIO(file_bytes))
        parts = []
        for p in doc.paragraphs:
            if p.text and p.text.strip():
                parts.append(p.text.strip())
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    t = (cell.text or "").strip()
                    if t:
                        parts.append(t)
        return "\n".join(parts)
    except Exception as e:
        logger.error(f"Docx Parse Error: {e}")
        return ""


# -------------------------
# Helper: JSON cleanup
# -------------------------
def _extract_json_text(raw: str) -> str:
    if not raw:
        return ""
    s = raw.strip()

    # strip code fences
    if s.startswith("```"):
        s = re.sub(r"^```[a-zA-Z0-9]*\s*", "", s)
        s = re.sub(r"\s*```$", "", s).strip()

    # keep the outermost {...}
    start = s.find("{")
    end = s.rfind("}")
    if start != -1 and end != -1 and end > start:
        s = s[start : end + 1].strip()
    return s


def _ensure_edge_ids(edges: list) -> None:
    for i, e in enumerate(edges, start=1):
        if isinstance(e, dict) and not e.get("id"):
            e["id"] = f"e{i}"


def _clamp_weight(edges: list) -> None:
    for e in edges:
        if not isinstance(e, dict):
            continue
        vis = e.get("visual")
        if not isinstance(vis, dict):
            vis = {}
            e["visual"] = vis

        w = vis.get("weight")
        try:
            w = float(w)
        except Exception:
            w = 0.5

        if w < 0:
            w = 0.0
        if w > 1:
            w = 1.0
        vis["weight"] = w


def _calculate_graph_metrics(graph: dict) -> dict:
    nodes = graph.get("nodes", [])
    edges = graph.get("edges", [])
    detected_language = graph.get("detected_language", "English")

    if not isinstance(nodes, list) or not isinstance(edges, list):
        raise ValueError("Graph JSON must contain 'nodes'(list) and 'edges'(list).")

    # 1. Prepare NetworkX Graph
    G = nx.Graph()

    # Normalize Node IDs
    node_map = {} # Map ID to Node Object
    for n in nodes:
        if isinstance(n, dict) and n.get("id") is not None:
            n["id"] = str(n["id"])
            
            # Bilingual fillers
            if "label_en" not in n and "label" in n:
                n["label_en"] = n["label"]
                n["label_original"] = n["label"]
            if "description_en" not in n and "description" in n:
                n["description_en"] = n["description"]
                n["description_original"] = n["description"]
            
            G.add_node(n["id"])
            node_map[n["id"]] = n

    # Normalize Edges & Build Graph
    valid_edges = []
    for e in edges:
        if not isinstance(e, dict):
            continue
        s = e.get("source")
        t = e.get("target")
        if s is None or t is None:
            continue

        s = str(s)
        t = str(t)
        e["source"] = s
        e["target"] = t

        if s in node_map and t in node_map:
            valid_edges.append(e)
            G.add_edge(s, t)

            # Bilingual fillers for edges
            rel = e.get("relation", {})
            if isinstance(rel, dict):
                if "label_en" not in rel and "label" in rel:
                    rel["label_en"] = rel["label"]
                    rel["label_original"] = rel["label"]
                if "description_en" not in rel and "description" in rel:
                    rel["description_en"] = rel["description"]
                    rel["description_original"] = rel["description"]
                if "type_en" not in rel and "type" in rel:
                    rel["type_en"] = rel["type"]
                    rel["type_original"] = rel["type"]

    # 2. Calculate Metrics using NetworkX
    
    # Degree (for Size)
    degrees = dict(G.degree())
    
    # Betweenness Centrality (for Color/Heatmap)
    # This detects "bridges" - nodes that connect different clusters
    try:
        betweenness = nx.betweenness_centrality(G)
    except Exception as e:
        logger.warning(f"Betweenness calculation failed: {e}")
        betweenness = {n: 0 for n in G.nodes()}

    # 3. Assign metrics back to nodes
    for n in nodes:
        if not isinstance(n, dict):
            continue
        nid = n.get("id")
        if nid in node_map:
            # We use Betweenness for the 'centrality' field now (Influence)
            n["centrality"] = float(betweenness.get(nid, 0.0))
            
            vis = n.get("visual")
            if not isinstance(vis, dict):
                vis = {}
                n["visual"] = vis
            
            # We use Degree for Visual Size (Popularity)
            d = degrees.get(nid, 0)
            vis["size"] = 15 + (d * 5)

    _ensure_edge_ids(valid_edges)
    _clamp_weight(valid_edges)

    return {
        "detected_language": detected_language,
        "total_characters": len([n for n in nodes if isinstance(n, dict)]),
        "nodes": nodes,
        "edges": valid_edges,
    }


def _is_timeline_schema(parsed: Any) -> bool:
    return isinstance(parsed, dict) and isinstance(parsed.get("timeline"), list)


def _process_model_output(parsed: dict) -> dict:
    if _is_timeline_schema(parsed):
        timeline = parsed.get("timeline", [])
        for i, phase in enumerate(timeline, start=1):
            if not isinstance(phase, dict):
                continue
            graph = phase.get("graph")
            if isinstance(graph, dict):
                phase["graph"] = _calculate_graph_metrics(graph)
            else:
                logger.warning(f"Phase {i} missing graph object; skipped metrics.")

        return {"timeline": timeline}

    return _calculate_graph_metrics(parsed)


# -------------------------
# Prompts
# -------------------------
SYSTEM_PROMPT = """
You are a graph data generator.

Your task is to extract the people and their direct relationships from the provided content (text or video) and output a JSON object with a timeline of phases. Each phase contains one graph.

### Output Constraints
1. Output ONLY valid JSON. No explanations, no markdown, no code fences.
2. Output must support bilingual display (Original Language + English).
3. You MUST output a top-level object with this structure:
{
  "timeline": [
    {
      "phase_id": "p1",
      "phase_name_original": "string",
      "phase_name_en": "string",
      "summary_original": "string",
      "summary_en": "string",
      "graph": {
        "detected_language": "string",
        "total_characters": number,
        "nodes": [
          {
            "id": "string",
            "label_original": "string",
            "label_en": "string",
            "description_original": "string",
            "description_en": "string"
          }
        ],
        "edges": [
          {
            "id": "string",
            "source": "node_id",
            "target": "node_id",
            "relation": {
              "type_original": "string",
              "type_en": "string",
              "label_original": "string",
              "label_en": "string",
              "description_original": "string",
              "description_en": "string"
            },
            "visual": {
              "color": "string",
              "weight": number
            }
          }
        ]
      }
    }
  ]
}

### Language Handling
4. "detected_language": The primary language of the input content. MUST be the full English name (e.g. "English", "French", "Chinese"), NOT a code like "en" or "zh".
5. `_original` fields: content in the source language.
6. `_en` fields: content translated to English.
7. `node.id` MUST be a language-neutral, lowercase snake_case identifier and must be stable across phases.

### Phase Rules
9. Divide the content into logical phases in chronological order.
10. Each phase has its own graph.
11. `phase_name` and `summary` MUST be bilingual.

### Node & Edge Definitions
12. Nodes represent people only.
13. Edges represent their direct relationships.
14. No Orphan Nodes: Every node in the graph SHOULD have at least one edge connecting it to the rest of the network, unless the character is explicitly described as having no contact with anyone.
15. Foundation First: Before linking interactions, establish the "Skeleton" (Parents, Siblings, Spouses).

### Identity Consistency
16. If the same character appears in multiple phases, reuse the EXACT same node.id.

### Visual Rules
17. edge.visual fields are required: color and weight (0.0 - 1.0).
18. Strict 1-to-1 color mapping for relationship types.
""".strip()

USER_PROMPT_TEMPLATE = """
Extract a timeline of character relationship graphs from the attached content.
If the content is a video, analyze the visual and audio narrative.

Return JSON ONLY.
""".strip()

CHAT_SYSTEM_PROMPT = """
You are the Nodal Graph Chat assistant.
You answer questions about characters, relationships, and plot based on the analyzed graph.
""".strip()


class ChatRequest(BaseModel):
    message: str
    context: Dict[str, Any]
    history: List[Dict[str, str]] = []


# -------------------------
# Streaming Logic
# -------------------------
def generate_analysis_stream(contents: List[types.Content], model: str, config: Dict[str, Any]) -> Generator[str, None, None]:
    """
    Generates a stream of NDJSON events:
    {"type": "thought", "text": "..."}
    {"type": "result", "data": {...}}
    """
    
    # Enable thinking config for Gemini 3
    # Use types.ThinkingConfig as per latest SDK patterns
    config["thinking_config"] = types.ThinkingConfig(include_thoughts=True)

    try:
        response_stream = client.models.generate_content_stream(
            model=model,
            contents=contents,
            config=config
        )

        full_json_response = ""

        for chunk in response_stream:
            # chunk is a GenerateContentResponse
            if not chunk.candidates:
                continue
            
            # Iterate through parts to separate thoughts from model text
            for part in chunk.candidates[0].content.parts:
                if part.thought:
                    # Stream thought chunk to client
                    event = {"type": "thought", "text": part.text}
                    yield json.dumps(event) + "\n"
                elif part.text:
                    # Accumulate JSON text for final parsing
                    full_json_response += part.text

        # After stream finishes, parse the accumulated JSON
        try:
            json_text = _extract_json_text(full_json_response)
            parsed_json = json.loads(json_text)
            final_data = _process_model_output(parsed_json)
            
            event = {"type": "result", "data": final_data}
            yield json.dumps(event) + "\n"
        except json.JSONDecodeError as e:
            logger.error(f"Failed to parse final JSON: {e}")
            logger.error(f"Raw output: {full_json_response[:1000]}...")
            yield json.dumps({"type": "error", "message": "Failed to parse AI response into valid JSON."}) + "\n"
        except Exception as e:
            logger.error(f"Processing error: {e}")
            yield json.dumps({"type": "error", "message": str(e)}) + "\n"

    except Exception as e:
        logger.error(f"Stream Error: {e}")
        yield json.dumps({"type": "error", "message": str(e)}) + "\n"


@app.post("/api/analyze")
async def analyze_content(
    file: UploadFile = File(None),
    text_content: str = Form(None),
    video_url: str = Form(None),
):
    if not client:
        raise HTTPException(status_code=500, detail="Gemini client not initialized. Check API_KEY.")

    temp_video_path = None
    uploaded_video_file = None

    try:
        content_parts = []
        
        # 1. Handle File Upload (Document or Video)
        if file:
            content_type = file.content_type or ""
            filename = file.filename.lower() if file.filename else ""
            
            # --- VIDEO FILE HANDLING (Direct Gemini Upload) ---
            if content_type.startswith("video/") or filename.endswith((".mp4", ".mov", ".mpeg")):
                logger.info(f"Processing uploaded video file: {file.filename}")
                
                # Save to temp file because client.files.upload needs a path
                with tempfile.NamedTemporaryFile(delete=False, suffix=".mp4") as tmp:
                    shutil.copyfileobj(file.file, tmp)
                    temp_video_path = tmp.name
                
                logger.info(f"Uploading {file.filename} to Gemini File API...")
                uploaded_video_file = client.files.upload(path=temp_video_path)
                
                # Poll for processing completion
                while uploaded_video_file.state.name == "PROCESSING":
                    logger.info("Video is processing by Gemini...")
                    time.sleep(2)
                    uploaded_video_file = client.files.get(name=uploaded_video_file.name)
                
                if uploaded_video_file.state.name == "FAILED":
                    raise HTTPException(status_code=400, detail="Video processing failed by Gemini.")
                
                logger.info("Video ready. Adding to request.")
                # Add file_data part for the uploaded video
                content_parts.append(types.Part(
                    file_data=types.FileData(
                        file_uri=uploaded_video_file.uri, 
                        mime_type=uploaded_video_file.mime_type
                    )
                ))
            
            # --- PDF HANDLING (Direct Gemini Input) ---
            elif content_type == "application/pdf" or filename.endswith(".pdf"):
                logger.info(f"Processing uploaded PDF file: {file.filename}")
                file_bytes = await file.read()

                # Local pre-check: reject PDFs exceeding the supported page limit (1000)
                try:
                    reader = PdfReader(io.BytesIO(file_bytes))
                    page_count = len(reader.pages)
                except Exception as e:
                    logger.error(f"Failed to read PDF page count: {e}")
                    page_count = None

                if page_count is not None and page_count > 1000:
                    raise HTTPException(
                        status_code=400,
                        detail=f"The document contains {page_count} pages which exceeds the supported page limit of 1000."
                    )

                content_parts.append(types.Part.from_bytes(
                    data=file_bytes,
                    mime_type="application/pdf"
                ))

            # --- TEXT/DOCX HANDLING ---
            else:
                file_bytes = await file.read()
                text = ""
                if "wordprocessingml" in content_type or filename.endswith(".docx"):
                    text = extract_text_from_docx(file_bytes)
                else:
                    text = file_bytes.decode("utf-8", errors="ignore")
                
                if text:
                     content_parts.append(types.Part(text=text))

        # 2. Handle Video URL (Direct Pass)
        elif video_url:
            logger.info(f"Using Video URL: {video_url}")
            # Directly pass the URL to Gemini as requested by user
            content_parts.append(types.Part(
                file_data=types.FileData(file_uri=video_url)
            ))

        # 3. Handle Raw Text
        elif text_content:
            content_parts.append(types.Part(text=text_content))

        # Add prompts
        final_contents = [
            types.Content(
                parts=[
                    types.Part(text=SYSTEM_PROMPT),
                    *content_parts,
                    types.Part(text=USER_PROMPT_TEMPLATE)
                ]
            )
        ]

        if not content_parts: # No content found
             raise HTTPException(status_code=400, detail="No valid content provided for analysis.")

        logger.info("Starting Gemini Stream...")
        
        # Return StreamingResponse for real-time thought updates
        return StreamingResponse(
            generate_analysis_stream(
                contents=final_contents,
                model="gemini-3-pro-preview",
                config={"response_mime_type": "application/json"}
            ),
            media_type="application/x-ndjson"
        )

    except Exception as e:
        logger.error(f"Analysis Setup Error: {e}")
        if temp_video_path and os.path.exists(temp_video_path):
            try:
                os.remove(temp_video_path)
            except Exception:
                pass
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/chat")
async def chat_with_context(request: ChatRequest):
    if not client:
        raise HTTPException(status_code=500, detail="Gemini client not initialized.")

    try:
        graph_context = request.context or {}
        history_str = ""
        if request.history:
            for msg in request.history[-10:]:
                role = "User" if msg.get("role") == "user" else "Assistant"
                history_str += f"{role}: {msg.get('content')}\n"

        prompt = (
            "CONTEXT (Graph JSON):\n"
            f"{json.dumps(graph_context, ensure_ascii=False)}\n\n"
            "CONVERSATION HISTORY:\n"
            f"{history_str}\n"
            "USER QUESTION:\n"
            f"{request.message}"
        )

        def generate_chat_stream():
            stream = client.models.generate_content_stream(
                model="gemini-3-pro-preview",
                contents=[types.Content(parts=[
                    types.Part(text=CHAT_SYSTEM_PROMPT), 
                    types.Part(text=prompt)
                ])],
            )
            for chunk in stream:
                if chunk.text:
                    yield chunk.text

        # Return plain text stream
        return StreamingResponse(generate_chat_stream(), media_type="text/plain")

    except Exception as e:
        logger.error(f"Chat Error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
