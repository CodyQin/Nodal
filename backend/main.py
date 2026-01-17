import os
import json
import io
import re
from collections import Counter
from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from dotenv import load_dotenv
from google import genai
from google.genai import types
from pypdf import PdfReader
from docx import Document

load_dotenv()

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

client = genai.Client(api_key=os.getenv("GEMINI_API_KEY"))

# --- 辅助函数：提取文本 ---
def extract_text_from_pdf(file_bytes: bytes) -> str:
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        text = ""
        for page in reader.pages:
            page_text = page.extract_text() or ""
            text += page_text + "\n"
        return text
    except Exception as e:
        print(f"PDF Parse Error: {e}")
        return ""

def extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        doc = Document(io.BytesIO(file_bytes))

        parts = []

        # Paragraphs
        for p in doc.paragraphs:
            if p.text and p.text.strip():
                parts.append(p.text.strip())

        # Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        parts.append(text)

        return "\n".join(parts)
    except Exception as e:
        print(f"Docx Parse Error: {e}")
        return ""



# ====== Prompts ======
SYSTEM_PROMPT = """
You are a graph data generator.

Your task is to extract all the people and their relationships from the input content and output a single JSON object that strictly follows the schema.

### Core Constraints
1. Output ONLY valid JSON. No explanations, no markdown.
2. All output content MUST be in English.
3. Total_characters must equal the number of nodes.

### Node & Edge Definitions
4. Nodes represent people only.
5. Degree is defined as the total number of direct relationships (edges) connected to a node.
6. Edges represent direct relationships only.
7. Each node must appear in at least one edge unless the text strongly implies isolation.

### Extraction Logic & Detail (CRITICAL)
8. Be conservative: Do not create relationships unless implied by the content.
9. Weight Calculation: Weight must be a number between 0 and 1. Calculate weight based on interaction intensity and emotional depth, not just social hierarchy or bureaucratic reporting.
10. Rich Relationship Descriptions:
    - The description field for each relationship is NOT optional.
    - It MUST be more than three sentences long.
    - Content Requirement: Do not simply restate the label. You must analyze the relationship by citing specific interactions, key events, and the emotional nuance found in the text to justify the connection.

### Visual & Normalization Rules
11. Relationship Normalization:
    - You MUST normalize relationship labels to a standard, lowercase format before assigning colors.
12. Strict Color Consistency (1-to-1 Mapping):
    - One Type = One Color; One Color = One Type.
13. Color Palette: Use distinct, high-contrast colors.
14. All visual fields are required.
""".strip()

USER_PROMPT_TEMPLATE = """
Extract a character relationship graph from the following content.

Output format must be exactly:

{
  "total_characters": number,
  "nodes": [
    {
      "id": "string",
      "label": "string",
      "description": "string"
    }
  ],
  "edges": [
    {
      "id": "string",
      "source": "node_id",
      "target": "node_id",
      "relation": {
        "type": "string",
        "label": "string",
        "description": "string"
      },
      "visual": {
        "color": "string",
        "weight": number
      }
    }
  ]
}

The content is provided in the attached file part.
""".strip()


def _ensure_edge_ids(edges: list) -> None:
    """If an edge has no id, fill sequential ids e1,e2,..."""
    for i, e in enumerate(edges, start=1):
        if isinstance(e, dict) and not e.get("id"):
            e["id"] = f"e{i}"

def _clamp_weight(edges: list) -> None:
    """Ensure edge.visual.weight is float in [0,1]."""
    for e in edges:
        if not isinstance(e, dict):
            continue
        vis = e.get("visual")
        if not isinstance(vis, dict):
            vis = {}
            e["visual"] = vis
        w = vis.get("weight")
        try:
            w = float(w)
        except Exception:
            w = 0.5
        if w < 0:
            w = 0.0
        if w > 1:
            w = 1.0
        vis["weight"] = w

def _extract_json_text(raw: str) -> str:
    """
    Extract a JSON object string from model output.
    Handles markdown fences like ```json ... ``` and extra text before/after.
    """
    if not raw:
        return ""

    s = raw.strip()

    # Remove markdown code fences if present
    if s.startswith("```"):
        # remove first fence line: ``` or ```json
        s = re.sub(r"^```[a-zA-Z0-9]*\s*", "", s)
        # remove trailing ```
        s = re.sub(r"\s*```$", "", s).strip()

    # If still has leading/trailing junk, grab the first {...} block
    start = s.find("{")
    end = s.rfind("}")
    if start != -1 and end != -1 and end > start:
        s = s[start:end+1].strip()

    return s

def _add_degree_and_size(parsed: dict) -> dict:
    """
    Add centrality (degree) and visual.size for each node.
    Also fixes total_characters to equal len(nodes).
    Drops edges referencing missing nodes to avoid frontend issues.
    """
    nodes = parsed.get("nodes", [])
    edges = parsed.get("edges", [])

    if not isinstance(nodes, list) or not isinstance(edges, list):
        raise ValueError("Model JSON must contain 'nodes' (list) and 'edges' (list).")

    # Build node id set
    node_ids = []
    for n in nodes:
        if isinstance(n, dict) and n.get("id") is not None:
            node_ids.append(str(n["id"]))
            n["id"] = str(n["id"])  # normalize to string
    node_set = set(node_ids)

    # Validate edges endpoints & compute degree
    degree = Counter()
    missing_endpoints = []

    valid_edges = []
    for e in edges:
        if not isinstance(e, dict):
            continue
        s = e.get("source")
        t = e.get("target")
        if s is None or t is None:
            continue

        s = str(s)
        t = str(t)
        e["source"] = s
        e["target"] = t

        if s not in node_set:
            missing_endpoints.append(("source", s, e.get("id")))
        if t not in node_set:
            missing_endpoints.append(("target", t, e.get("id")))

        if s in node_set and t in node_set:
            valid_edges.append(e)
            degree[s] += 1
            degree[t] += 1

    if missing_endpoints:
        print("⚠️ Missing endpoints found (edges dropped):")
        for item in missing_endpoints[:20]:
            print("  ", item)
        if len(missing_endpoints) > 20:
            print(f"  ... and {len(missing_endpoints) - 20} more")

    # Ensure each node has centrality & visual.size
    for n in nodes:
        if not isinstance(n, dict):
            continue
        nid = str(n.get("id", ""))
        d = int(degree.get(nid, 0))
        n["centrality"] = d

        vis = n.get("visual")
        if not isinstance(vis, dict):
            vis = {}
            n["visual"] = vis
        vis["size"] = d * 6 + 10

        if "description" not in n and "desc" in n:
            n["description"] = n.pop("desc")

    _ensure_edge_ids(valid_edges)
    _clamp_weight(valid_edges)

    return {
        "total_characters": len([n for n in nodes if isinstance(n, dict)]),
        "nodes": nodes,
        "edges": valid_edges
    }


@app.post("/api/analyze")
async def analyze_content(
    file: UploadFile = File(None),
    text_content: str = Form(None)
):
    try:
        final_text = ""

        # 1) 优先处理文件上传
        if file:
            print(f"Processing file: {file.filename} ({file.content_type})")
            content = await file.read()

            if file.content_type == "application/pdf":
                final_text = extract_text_from_pdf(content)
            elif "wordprocessingml" in file.content_type or file.filename.endswith(".docx"):
                final_text = extract_text_from_docx(content)
            elif file.content_type.startswith("text/"):
                final_text = content.decode("utf-8", errors="ignore")
            else:
                try:
                    final_text = content.decode("utf-8", errors="ignore")
                except Exception:
                    raise HTTPException(status_code=400, detail="Unsupported file format")

        # 2) 如果没有文件，检查是否有粘贴的文本
        elif text_content:
            final_text = text_content

        if not final_text or not final_text.strip():
            raise HTTPException(status_code=400, detail="No content provided")

        print(f"Extracted text length: {len(final_text)} characters")
        print(f"Extracted text preview: {final_text[:500]}...")

        # 3) prompt + 上传文件（以 text/plain file-part 形式传给模型）
        user_prompt = USER_PROMPT_TEMPLATE

        file_part = types.Part.from_bytes(
            data=final_text.encode("utf-8", errors="ignore"),
            mime_type="text/plain"
        )

        response = client.models.generate_content(
            model="gemini-3-pro-preview",
            contents=[SYSTEM_PROMPT, user_prompt, file_part],
        )

        # 4) 解析模型 JSON
        raw = response.text
        json_text = _extract_json_text(raw)
        parsed = json.loads(json_text)

        # 5) 补 centrality(degree) & visual.size，并修正 total_characters
        final_graph = _add_degree_and_size(parsed)

        # Optional: save latest result to a file for debugging
        try:
            with open("latest_graph.json", "w", encoding="utf-8") as f:
                json.dump(final_graph, f, ensure_ascii=False, indent=2)
            print("Saved to latest_graph.json")
        except Exception as e:
            print("Save JSON failed:", e)

       # print("Generated final graph result:")
       # print(json.dumps(final_graph, indent=2, ensure_ascii=False))

        return final_graph

    except json.JSONDecodeError as je:
        print(f"JSON Decode Error: {je}")
        print(f"Raw response: {raw if 'raw' in locals() else 'Response not available'}")
        raise HTTPException(status_code=500, detail=f"Invalid JSON response from AI: {str(je)}")
    except Exception as e:
        print(f"Error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
